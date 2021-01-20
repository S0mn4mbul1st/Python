import re

from PIL import Image
from pip._vendor import requests
from docx import Document
from docx.shared import Inches

first_chapter = []
book_title = ""
author = ""
book_path = 'Book.txt'


def get_author_name():  # 2
    f = open(book_path, encoding='utf-8')
    lines = f.read()
    a = re.compile("(Author:)(.*)")
    author = (a.findall(lines))[0][1]

    return author


def get_title():  # 2
    f = open(book_path, encoding='utf-8')
    lines = f.read()
    f = re.compile("(Title:)(.*)")
    book_title = (f.findall(lines))[0][1]

    return book_title


def get_first_chapter():
    txt = ""
    try:
        f = open(book_path, encoding='utf-8')
        file_line = f.readlines()
        for line in file_line:
            txt += line
            if line.startswith("the prince himself, or of others, or else by fortune or by ability."):
                break
        start_Index = txt.find('CHAPTER I --')
        myChapter = txt[start_Index:]

    except FileNotFoundError:
        return ("file not found")

    return myChapter


def png_downloader(url):
    response = requests.get(url)

    file = open("#1.png", "wb")
    file.write(response.content)
    file.close()


def png_cropper():
    im = Image.open(r"#1.png")
    width, height = im.size
    left = 5
    top = height / 4
    right = 164
    bottom = 3 * height / 4
    im1 = im.crop((left, top, right, bottom))
    im1.show()
    im1.save('#1-cropped.jpg')


def png_rotated_paster():
    im = Image.open('#1.png')
    colorImage = Image.open('#2.jpg')
    rotated = colorImage.rotate(45)
    rotated.save('#2-rotated.jpg')

    im.paste(rotated)
    im.show()
    im.save('#1-logo-pasted.jpg')


def word_document():
    document = Document()
    document.add_heading(get_title(), 0)
    p = document.add_paragraph()
    p.add_run('\n' + 'Author:' + get_author_name()).bold = True
    p.add_run('\n' + 'Title:' + get_title()).bold = True
    p.add_run("\n\n")
    p.add_run('\n' + 'Aykhan Imanov\n').bold = True
    p.add_run(get_first_chapter()).bold = False
    document.add_picture('#1-logo-pasted.jpg', Inches(5), Inches(5))

    document.add_page_break()
    document.save('summary.docx')

def test_package():

    png_downloader("https://images-na.ssl-images-amazon.com/images/I/71de1OC-ZIL.jpg")
    png_cropper()
    png_rotated_paster()
    word_document()


test_package()
